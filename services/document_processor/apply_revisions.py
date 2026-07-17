"""Apply approved review findings into a source .docx body (new edition)."""

from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt

from services.ai_orchestrator.review_findings import normalize_revision_action
from services.document_processor.annotated_export import (
    _add_plain_heading,
    _find_paragraph,
)

# Leading clause number in paragraph body, e.g. "1.2.", "1.2", "п. 1.2.", "Пункт 1.2 "
_LEADING_NUMBER_RE = re.compile(
    r"^\s*((?:п(?:ункт)?\.?\s*)?\d+(?:\.\d+)*\.?)\s*",
    re.IGNORECASE,
)
# Number extracted from clause_ref like "П. 1.2", "п.1.2.", "1.2"
_CLAUSE_REF_NUMBER_RE = re.compile(
    r"(?:п(?:ункт)?\.?\s*)?(\d+(?:\.\d+)*)\.?",
    re.IGNORECASE,
)


def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text, keeping formatting of the first run when possible."""
    text = text or ""
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
        return
    paragraph.add_run(text)


def _has_word_auto_numbering(paragraph) -> bool:
    """True when Word list numbering (numPr) is attached to the paragraph."""
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    return p_pr.numPr is not None


def _extract_leading_number(text: str) -> tuple[str, str]:
    """Return (leading_number_token, remainder). Token without trailing spaces."""
    match = _LEADING_NUMBER_RE.match(text or "")
    if not match:
        return "", (text or "").strip()
    return match.group(1).strip(), (text or "")[match.end() :].strip()


def _number_from_clause_ref(clause_ref: str) -> str:
    """Normalize clause_ref into a paragraph number like '1.2.'."""
    raw = (clause_ref or "").strip()
    if not raw:
        return ""
    match = _CLAUSE_REF_NUMBER_RE.search(raw)
    if not match:
        return ""
    return f"{match.group(1)}."


def _suggested_has_number(suggested: str) -> bool:
    token, _ = _extract_leading_number(suggested)
    if token:
        return True
    return bool(re.match(r"^\s*\d+(?:\.\d+)*\.?\s*", suggested or ""))


def _preserve_clause_number(paragraph, suggested: str, clause_ref: str) -> str:
    """Ensure restate text keeps the visible clause number from the source paragraph.

    If Word auto-numbering is used, the number lives outside runs — leave text as-is.
    Otherwise take the leading number from the paragraph body, or fall back to clause_ref.
    """
    suggested = (suggested or "").strip()
    if not suggested:
        return suggested
    if _has_word_auto_numbering(paragraph):
        return suggested
    if _suggested_has_number(suggested):
        return suggested

    existing = paragraph.text or ""
    prefix, _ = _extract_leading_number(existing)
    if not prefix:
        prefix = _number_from_clause_ref(clause_ref)
    if not prefix:
        return suggested

    if not prefix.endswith("."):
        prefix = f"{prefix}."
    return f"{prefix} {suggested.lstrip()}"


def _apply_to_paragraph(paragraph, finding: dict) -> None:
    suggested = (finding.get("suggested_revision") or "").strip()
    if not suggested:
        return
    action = normalize_revision_action(finding)
    clause_ref = (finding.get("clause_ref") or "").strip()
    if action == "supplement":
        existing = (paragraph.text or "").rstrip()
        joined = f"{existing} {suggested}".strip() if existing else suggested
        _set_paragraph_text(paragraph, joined)
    else:
        numbered = _preserve_clause_number(paragraph, suggested, clause_ref)
        _set_paragraph_text(paragraph, numbered)


def apply_revisions_to_docx(
    source: Path | str | bytes,
    findings: list[dict],
) -> tuple[bytes, dict]:
    """Return revised .docx bytes and stats: {applied, unmatched, skipped}.

    - restate: replace matched paragraph with suggested_revision (clause number preserved)
    - supplement: append suggested_revision to matched paragraph
    Unmatched findings go to an appendix at the end of the document.
    """
    if isinstance(source, (bytes, bytearray)):
        doc = DocxDocument(BytesIO(source))
    else:
        doc = DocxDocument(str(source))

    used: set[int] = set()
    applied: list[dict] = []
    unmatched: list[dict] = []
    skipped = 0

    for finding in findings:
        if not isinstance(finding, dict):
            skipped += 1
            continue
        suggested = (finding.get("suggested_revision") or "").strip()
        if not suggested:
            skipped += 1
            continue

        quote = (finding.get("original_text") or "").strip()
        para = _find_paragraph(doc, quote, used) if quote else None
        if para is None:
            clause = (finding.get("clause_ref") or "").strip()
            if clause and len(clause) >= 3:
                para = _find_paragraph(doc, clause, used)
        if para is None:
            unmatched.append(finding)
            continue

        _apply_to_paragraph(para, finding)
        applied.append(finding)

    if unmatched:
        doc.add_page_break()
        _add_plain_heading(doc, "Правки без привязки к тексту")
        doc.add_paragraph(
            "Ниже — правки, которые не удалось однозначно вставить в исходный документ. "
            "Внесите их вручную или уточните цитату в проверке."
        )
        for i, f in enumerate(unmatched, start=1):
            clause = (f.get("clause_ref") or f"п. {i}").strip()
            action = normalize_revision_action(f)
            label = "Дополнить" if action == "supplement" else "Изложить в редакции"
            text = (f.get("suggested_revision") or "").strip()
            if action != "supplement" and not _suggested_has_number(text):
                num = _number_from_clause_ref(clause)
                if num:
                    text = f"{num} {text}".strip()
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {clause} — {label}: ")
            run.bold = True
            p.add_run(text)
            orig = (f.get("original_text") or "").strip()
            if orig:
                note = doc.add_paragraph()
                r = note.add_run(f"Исходная цитата: «{orig[:400]}»")
                r.italic = True
                r.font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), {
        "applied": len(applied),
        "unmatched": len(unmatched),
        "skipped": skipped,
    }
