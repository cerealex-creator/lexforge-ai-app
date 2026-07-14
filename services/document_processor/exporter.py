"""Build a Word (.docx) report from a completed contract review task."""

from datetime import datetime
from io import BytesIO

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SEVERITY_LABELS = {
    "critical": "Критично",
    "high": "Высокая",
    "medium": "Средняя",
    "low": "Низкая",
}
SEVERITY_COLORS = {
    "critical": "F8CBAD",
    "high": "FCE4D6",
    "medium": "FFF2CC",
    "low": "E2EFDA",
}

MODE_LABELS = {"full": "Полная проверка", "errors": "Ошибки", "risks": "Угрозы и риски"}
INDUSTRY_LABELS = {
    "construction": "Строительство",
    "production": "Производство",
    "supply": "Поставки",
    "general": "Универсальное",
}


def _shade_cell(cell, color_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _bold_cell(cell, text: str) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True


def _risk_color(score: int | None) -> str:
    if score is None:
        return "D9D9D9"
    if score >= 9:
        return "C00000"
    if score >= 7:
        return "ED7D31"
    if score >= 4:
        return "FFC000"
    return "70AD47"


def build_review_report(
    *,
    document_title: str,
    company_name: str,
    review_mode: str,
    industry: str,
    user_comment: str | None,
    completed_at: datetime | None,
    risk_score: int | None,
    risk_rationale: str | None,
    findings: list[dict],
) -> bytes:
    doc = DocxDocument()

    doc.add_heading("Заключение по результатам проверки договора", level=1)

    meta_lines = [
        ("Документ", document_title),
        ("Компания", company_name),
        ("Режим проверки", MODE_LABELS.get(review_mode, review_mode)),
        ("Отрасль", INDUSTRY_LABELS.get(industry, industry)),
        ("Дата проверки", completed_at.strftime("%d.%m.%Y %H:%M") if completed_at else "—"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    if user_comment and user_comment.strip():
        p = doc.add_paragraph()
        p.add_run("Комментарий юриста: ").bold = True
        p.add_run(user_comment.strip())

    doc.add_paragraph()

    risk_heading = doc.add_paragraph()
    risk_run = risk_heading.add_run(f"Оценка риска: {risk_score if risk_score is not None else '—'}/10")
    risk_run.bold = True
    risk_run.font.size = Pt(14)
    color = _risk_color(risk_score)
    risk_run.font.color.rgb = RGBColor.from_string(color)

    if risk_rationale:
        doc.add_paragraph(risk_rationale)

    doc.add_paragraph()
    doc.add_heading(f"Замечания ({len(findings)})", level=2)

    if not findings:
        doc.add_paragraph("Критических замечаний не выявлено.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["№", "Пункт", "Критичность", "Цитата / проблема", "Рекомендация"]
        for cell, text in zip(table.rows[0].cells, headers):
            _bold_cell(cell, text)
            _shade_cell(cell, "D9E2F3")

        for i, f in enumerate(findings, start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = f.get("clause_ref") or "—"

            severity = f.get("severity", "medium")
            row[2].text = SEVERITY_LABELS.get(severity, severity)
            _shade_cell(row[2], SEVERITY_COLORS.get(severity, "FFFFFF"))

            original = (f.get("original_text") or "").strip()
            rationale = (f.get("rationale") or "").strip()
            cell3 = row[3]
            cell3.text = ""
            if original:
                r = cell3.paragraphs[0].add_run(f"«{original}»")
                r.italic = True
            if rationale:
                p2 = cell3.add_paragraph() if original else cell3.paragraphs[0]
                p2.add_run(rationale)

            revision = (f.get("suggested_revision") or "").strip()
            if revision:
                from services.ai_orchestrator.review_findings import normalize_revision_action

                action = normalize_revision_action(f)
                prefix = "Дополнить: " if action == "supplement" else "Изложить в редакции: "
                row[4].text = prefix + revision
            else:
                row[4].text = "—"

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run("Документ сформирован автоматически LexForge AI. Требует проверки юристом перед использованием.")
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("808080")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


IMPACT_LABELS = {
    "favorable": "Выгодно нам",
    "unfavorable": "Невыгодно нам",
    "neutral": "Нейтрально",
    "suspicious": "Подозрительно",
}


def build_comparison_report(
    *,
    base_document_title: str,
    revised_document_title: str,
    company_name: str,
    user_comment: str | None,
    completed_at: datetime | None,
    risk_delta: int | None,
    summary: str | None,
    changes: list[dict],
) -> bytes:
    doc = DocxDocument()

    doc.add_heading("Заключение по сравнению редакций договора", level=1)

    meta_lines = [
        ("Базовая редакция", base_document_title),
        ("Новая редакция", revised_document_title),
        ("Компания", company_name),
        ("Дата сравнения", completed_at.strftime("%d.%m.%Y %H:%M") if completed_at else "—"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    if user_comment and user_comment.strip():
        p = doc.add_paragraph()
        p.add_run("Комментарий юриста: ").bold = True
        p.add_run(user_comment.strip())

    doc.add_paragraph()

    delta_heading = doc.add_paragraph()
    delta_text = "—"
    if risk_delta is not None:
        delta_text = f"{risk_delta:+d}" if risk_delta != 0 else "0"
    delta_run = delta_heading.add_run(f"Изменение риска: {delta_text}")
    delta_run.bold = True
    delta_run.font.size = Pt(14)

    if summary:
        doc.add_paragraph(summary)

    doc.add_paragraph()
    doc.add_heading(f"Изменения ({len(changes)})", level=2)

    if not changes:
        doc.add_paragraph("Различий между редакциями не обнаружено.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["№", "Пункт", "Влияние", "Критичность", "Было / Стало", "Обоснование"]
        for cell, text in zip(table.rows[0].cells, headers):
            _bold_cell(cell, text)
            _shade_cell(cell, "D9E2F3")

        for i, c in enumerate(changes, start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = c.get("clause_ref") or "—"
            row[2].text = IMPACT_LABELS.get(c.get("impact", ""), c.get("impact", "—"))

            severity = c.get("severity", "medium")
            row[3].text = SEVERITY_LABELS.get(severity, severity)
            _shade_cell(row[3], SEVERITY_COLORS.get(severity, "FFFFFF"))

            original = (c.get("original_text") or "").strip()
            revised = (c.get("revised_text") or "").strip()
            cell4 = row[4]
            cell4.text = ""
            if original:
                p = cell4.paragraphs[0]
                p.add_run("Было: ").bold = True
                p.add_run(original)
            if revised:
                p2 = cell4.add_paragraph() if original else cell4.paragraphs[0]
                p2.add_run("Стало: ").bold = True
                p2.add_run(revised)

            row[5].text = c.get("rationale") or "—"

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Документ сформирован автоматически LexForge AI. Требует проверки юристом перед использованием."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("808080")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Short role labels for protocol party headers (review_position → our side name)
_POSITION_ROLE_LABELS = {
    "contractor": "Подрядчик",
    "general_contractor": "Генподрядчик",
    "gc_vs_contractor": "Генподрядчик",
    "gc_vs_customer": "Генподрядчик",
    "customer": "Заказчик",
    "supplier": "Поставщик",
    "buyer": "Покупатель",
    "executor": "Поставщик услуг",
}


def _set_landscape(doc: DocxDocument) -> None:
    section = doc.sections[0]
    section.orientation = 1  # WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def build_disagreement_protocol(
    *,
    document_title: str,
    company_name: str,
    completed_at: datetime | None,
    findings: list[dict],
    our_party_label: str | None = None,
    their_party_label: str | None = None,
    review_position: str | None = None,
    include_our_comments: bool = True,
) -> bytes:
    """Tabular протокол разногласий: две редакции + комментарии каждой стороны."""
    our = (our_party_label or "").strip()
    if not our:
        our = _POSITION_ROLE_LABELS.get((review_position or "").strip(), "") or company_name or "Наша сторона"
    their = (their_party_label or "").strip() or "Контрагент"

    doc = DocxDocument()
    _set_landscape(doc)

    doc.add_heading("Протокол разногласий", level=1)

    meta_lines = [
        ("К договору", document_title),
        ("Сторона, направившая протокол", f"{our} ({company_name})" if company_name else our),
        ("Дата", completed_at.strftime("%d.%m.%Y") if completed_at else "—"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value))

    intro = doc.add_paragraph()
    intro.add_run(
        "Ниже изложены разногласия по отдельным положениям договора: редакция одной стороны, "
        "предлагаемая редакция другой стороны и поля комментариев для согласования."
    )

    doc.add_paragraph()

    if not findings:
        doc.add_paragraph("Разногласий для включения в протокол не выбрано.")
    else:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = [
            "№",
            "Пункт",
            f"Редакция ({their})",
            f"Редакция ({our})",
            f"Комментарий ({their})",
            f"Комментарий ({our})",
        ]
        for cell, text in zip(table.rows[0].cells, headers):
            _bold_cell(cell, text)
            _shade_cell(cell, "D9E2F3")

        for i, f in enumerate(findings, start=1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = (f.get("clause_ref") or "").strip() or "—"

            their_text = (f.get("original_text") or "").strip()
            our_text = (f.get("suggested_revision") or "").strip()
            row[2].text = their_text or "—"
            if our_text:
                from services.ai_orchestrator.review_findings import normalize_revision_action

                action = normalize_revision_action(f)
                prefix = "Дополнить: " if action == "supplement" else "Изложить в редакции: "
                row[3].text = prefix + our_text
            else:
                row[3].text = "—"

            # Their comment left blank for counterparty reply when printing/signing
            row[4].text = ""

            our_comment = ""
            if include_our_comments:
                our_comment = (f.get("lawyer_note") or "").strip() or (f.get("rationale") or "").strip()
            row[5].text = our_comment

    doc.add_paragraph()
    doc.add_paragraph(
        "Настоящий протокол разногласий составлен в двух экземплярах, имеющих одинаковую юридическую силу, "
        "по одному для каждой из сторон."
    )

    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    sig.cell(0, 0).text = f"{their}:"
    sig.cell(0, 1).text = f"{our}:"
    sig.cell(1, 0).text = "\n_________________ / _________________\nдолжность, ФИО, подпись, дата"
    sig.cell(1, 1).text = "\n_________________ / _________________\nдолжность, ФИО, подпись, дата"

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Черновик сформирован LexForge AI на основе проверки договора. Перед направлением контрагенту "
        "проверьте редакции и комментарии."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("808080")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
