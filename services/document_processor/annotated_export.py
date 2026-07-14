"""Annotate a source .docx with Word review comments from AI findings."""

from __future__ import annotations

import re
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import Pt
from lxml import etree

from services.document_processor.exporter import SEVERITY_LABELS

COMMENTS_RELTYPE = RT.COMMENTS
COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_PARTNAME = PackURI("/word/comments.xml")

EMPTY_COMMENTS_XML = (
    b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    b' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
)

DEFAULT_AUTHOR = "Юрист компании"


def author_initials(author: str) -> str:
    words = [w for w in re.split(r"\s+", (author or "").strip()) if w]
    if not words:
        return "ЮР"
    letters = []
    for w in words[:3]:
        ch = next((c for c in w if c.isalnum()), "")
        if ch:
            letters.append(ch.upper())
    return "".join(letters)[:3] or "ЮР"


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


DEFAULT_AUTHOR = "Юрист компании"
AI_DISCLAIMER = "(сформировано LexForge AI, требует проверки)"

ISSUE_TYPE_LABELS = {
    "errors": "ошибки",
    "risks": "риски",
    "financial": "финансы",
    "compliance": "compliance",
    "cascade_gap": "каскадный разрыв",
}


def _comment_body(
    finding: dict,
    index: int,
    author: str,
    *,
    include_metadata: bool = False,
    include_ai_disclaimer: bool = False,
) -> str | None:
    """Counterparty-facing comment: only the proposed wording (no internal risk rationale).

    Returns None if there is no suggested_revision — such findings are skipped in annotated export.
    Distinguishes «дополнить» vs «изложить в редакции» via revision_action.
    """
    from services.ai_orchestrator.review_findings import (
        normalize_revision_action,
        revision_proposal_phrase,
    )

    clause = (finding.get("clause_ref") or f"п. {index}").strip()
    suggested = (finding.get("suggested_revision") or "").strip()
    if not suggested:
        return None

    action = normalize_revision_action(finding)
    proposal = revision_proposal_phrase(clause, suggested, action)

    lines: list[str] = []
    if include_metadata:
        severity = finding.get("severity") or "medium"
        sev_label = SEVERITY_LABELS.get(severity, severity)
        lines.append(f"[{sev_label}] {clause}")
        issue_type = (finding.get("issue_type") or "").strip()
        if issue_type:
            type_label = ISSUE_TYPE_LABELS.get(issue_type, issue_type)
            lines.append(f"Тип: {type_label}")
        action_label = "дополнить" if action == "supplement" else "изложить в новой редакции"
        lines.append(f"Вид правки: {action_label}")

    lines.append(proposal)

    if include_ai_disclaimer:
        lines.append(f"— {author} {AI_DISCLAIMER}")

    return "\n".join(lines)


def _iter_paragraphs(doc: DocxDocument):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _find_paragraph(doc: DocxDocument, quote: str, used: set[int]):
    """Return best unused paragraph match, or None."""
    needle = _normalize(quote)
    if not needle or len(needle) < 8:
        return None

    candidates: list[tuple[int, int, object]] = []
    for p in _iter_paragraphs(doc):
        pid = id(p._p)
        if pid in used:
            continue
        hay = _normalize(p.text)
        if not hay:
            continue
        if needle in hay:
            score = 1000 + min(len(needle), 500) - abs(len(hay) - len(needle))
            candidates.append((score, pid, p))
            continue
        soft = needle[:120]
        if len(soft) >= 24 and soft in hay:
            score = 500 + len(soft)
            candidates.append((score, pid, p))

    if not candidates:
        return None
    candidates.sort(key=lambda x: x[0], reverse=True)
    _, pid, para = candidates[0]
    used.add(pid)
    return para


def _get_or_create_comments_part(document: DocxDocument) -> Part:
    try:
        return document.part.part_related_by(COMMENTS_RELTYPE)
    except KeyError:
        part = Part(COMMENTS_PARTNAME, COMMENTS_CONTENT_TYPE, EMPTY_COMMENTS_XML, document.part.package)
        document.part.relate_to(part, COMMENTS_RELTYPE)
        return part


def _next_comment_id(comments_elm) -> int:
    ids = [
        int(c.get(qn("w:id")))
        for c in comments_elm.findall(qn("w:comment"))
        if c.get(qn("w:id")) is not None
    ]
    return (max(ids) + 1) if ids else 0


def _add_comment_element(
    comments_elm,
    comment_id: int,
    text: str,
    created: datetime,
    *,
    author: str,
    initials: str,
) -> None:
    comment = OxmlElement("w:comment")
    comment.set(qn("w:id"), str(comment_id))
    comment.set(qn("w:author"), author)
    comment.set(qn("w:initials"), initials)
    comment.set(qn("w:date"), created.strftime("%Y-%m-%dT%H:%M:%SZ"))

    for line in text.split("\n"):
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = line
        r.append(t)
        p.append(r)
        comment.append(p)

    comments_elm.append(comment)


def _attach_comment_to_paragraph(paragraph, comment_id: int) -> None:
    p = paragraph._p
    start = OxmlElement("w:commentRangeStart")
    start.set(qn("w:id"), str(comment_id))
    p.insert(0, start)

    end = OxmlElement("w:commentRangeEnd")
    end.set(qn("w:id"), str(comment_id))
    p.append(end)

    ref_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "CommentReference")
    rPr.append(rStyle)
    ref_run.append(rPr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)
    p.append(ref_run)

    for run in paragraph.runs:
        try:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        except Exception:
            pass


def _add_plain_heading(doc: DocxDocument, text: str) -> None:
    """Add a bold paragraph without relying on built-in Heading styles.

    Some uploaded contracts strip or customize styles; python-docx's add_heading
    then fails with KeyError / ValueError: no style with name 'Heading 1'.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def annotate_docx_with_comments(
    source: Path | str | bytes,
    findings: list[dict],
    *,
    author: str | None = None,
    include_metadata: bool = False,
    include_ai_disclaimer: bool = False,
) -> tuple[bytes, dict]:
    """Return annotated docx bytes and stats: {matched, unmatched}."""
    author_name = (author or "").strip() or DEFAULT_AUTHOR
    initials = author_initials(author_name)

    if isinstance(source, (bytes, bytearray)):
        doc = DocxDocument(BytesIO(source))
    else:
        doc = DocxDocument(str(source))

    comments_part = _get_or_create_comments_part(doc)
    comments_elm = parse_xml(comments_part.blob)
    next_id = _next_comment_id(comments_elm)
    created = datetime.now(timezone.utc)
    used: set[int] = set()

    matched: list[dict] = []
    unmatched: list[dict] = []

    for i, finding in enumerate(findings, start=1):
        body = _comment_body(
            finding,
            i,
            author_name,
            include_metadata=include_metadata,
            include_ai_disclaimer=include_ai_disclaimer,
        )
        if body is None:
            continue

        quote = (finding.get("original_text") or "").strip()
        para = _find_paragraph(doc, quote, used) if quote else None
        if para is None:
            clause = (finding.get("clause_ref") or "").strip()
            if clause and len(clause) >= 3:
                para = _find_paragraph(doc, clause, used)
        if para is None:
            unmatched.append(finding)
            continue

        _add_comment_element(
            comments_elm,
            next_id,
            body,
            created,
            author=author_name,
            initials=initials,
        )
        _attach_comment_to_paragraph(para, next_id)
        matched.append(finding)
        next_id += 1

    if unmatched:
        doc.add_page_break()
        _add_plain_heading(doc, "Предложения без привязки к тексту")
        doc.add_paragraph(
            "Ниже — правки, для которых не удалось однозначно найти цитату в исходном .docx."
        )
        for i, f in enumerate(unmatched, start=1):
            proposal = _comment_body(
                f,
                i,
                author_name,
                include_metadata=include_metadata,
                include_ai_disclaimer=include_ai_disclaimer,
            )
            if not proposal:
                continue
            p = doc.add_paragraph()
            run = p.add_run(proposal)
            run.bold = False

    comments_part._blob = etree.tostring(comments_elm, xml_declaration=True, encoding="UTF-8", standalone=True)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), {"matched": len(matched), "unmatched": len(unmatched)}
