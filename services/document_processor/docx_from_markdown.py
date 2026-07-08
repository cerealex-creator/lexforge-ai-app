from __future__ import annotations

import re
from docx import Document as DocxDocument
from docx.shared import Pt


def markdown_to_docx_bytes(md: str) -> bytes:
    doc = DocxDocument()
    text = (md or "").strip()
    if not text:
        doc.add_paragraph("")
    else:
        for line in text.splitlines():
            l = line.strip()
            if not l:
                doc.add_paragraph("")
                continue
            if l.startswith("### "):
                doc.add_heading(l[4:].strip(), level=3)
            elif l.startswith("## "):
                doc.add_heading(l[3:].strip(), level=2)
            elif l.startswith("# "):
                doc.add_heading(l[2:].strip(), level=1)
            elif re.match(r"^\\d+\\.\\s+\\S+", l):
                # Numbered clause -> use a normal paragraph, but bold the number.
                p = doc.add_paragraph()
                m = re.match(r"^(\\d+\\.)\\s+(.*)$", l)
                if m:
                    run_num = p.add_run(m.group(1) + " ")
                    run_num.bold = True
                    p.add_run(m.group(2))
                else:
                    p.add_run(line)
            elif l.startswith("- ") or l.startswith("• "):
                doc.add_paragraph(l[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(line)

    from io import BytesIO

    # Improve default font size readability
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

