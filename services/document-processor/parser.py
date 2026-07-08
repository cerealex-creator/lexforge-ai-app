"""Parse docx and pdf into plain text."""

import hashlib
from pathlib import Path


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def parse_document(path: Path, mime_type: str) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx" or "wordprocessingml" in mime_type:
        return _parse_docx(path)
    if suffix == ".pdf" or mime_type == "application/pdf":
        return _parse_pdf(path)
    if suffix == ".txt" or mime_type.startswith("text/"):
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Неподдерживаемый формат: {suffix or mime_type}")


def _parse_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _parse_pdf(path: Path) -> str:
    import fitz

    doc = fitz.open(str(path))
    parts: list[str] = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            parts.append(text)
    doc.close()
    return "\n\n".join(parts)
